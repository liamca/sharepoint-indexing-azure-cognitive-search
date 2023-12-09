import io
import json
import os
from datetime import datetime, timedelta, timezone
from typing import Dict, List, Optional

import jwt
import msal
import requests
from docx import Document as DocxDocument
from langchain.docstore.document import Document
from dotenv import load_dotenv
from msal.application import ConfidentialClientApplication
from gbb_ai.azure_search_security_trimming import SecurityGroupManager

# load logging
from utils.ml_logging import get_logger

logger = get_logger()
manager_security = SecurityGroupManager()

# Load .env file
load_dotenv()

# Get the variables
tenantID = os.getenv("TENANT_ID")
clientID = os.getenv("CLIENT_ID")
clientSecret = os.getenv("CLIENT_SECRET")
graphURI = "https://graph.microsoft.com"
authority = "https://login.microsoftonline.com/" + tenantID
scope = ["https://graph.microsoft.com/.default"]

class CustomDocument:
    def __init__(self, page_content, metadata=None):
        self.page_content = page_content
        self.metadata = metadata


def msgraph_auth(
    client_id: str, client_secret: str, authority: str, scope: list
) -> ConfidentialClientApplication:
    """
    Authenticate with Microsoft Graph using MSAL for Python.

    This function authenticates with Microsoft Graph using the MSAL library. It attempts to acquire a token silently,
    and if that fails, it tries to acquire a new token. It decodes and prints the acquired access token and its expiry time.

    :param client_id: The application (client) ID of your Azure AD app registration.
    :param client_secret: The client secret for your Azure AD app registration.
    :param authority: The authority URL for your Azure AD tenant.
    :param scope: Scopes required for the token.
    :return: A dictionary with the access token and its expiry time if successful, None otherwise.
    :raises Exception: If there are issues in acquiring or decoding the token.
    """
    app = msal.ConfidentialClientApplication(
        client_id, authority=authority, client_credential=client_secret
    )

    try:
        access_token = app.acquire_token_silent(scope, account=None)
        if not access_token:
            access_token = app.acquire_token_for_client(scopes=scope)
            if "access_token" in access_token:
                logger.info("New access token retrieved....")
            else:
                logger.error(
                    "Error acquiring authorization token. Check your tenantID, clientID, and clientSecret."
                )
                return None
        else:
            logger.info("Token retrieved from MSAL Cache....")

        algorithms = ["RS256"]
        decoded_access_token = jwt.decode(
            access_token["access_token"],
            algorithms=algorithms,
            options={"verify_signature": False},
        )
        access_token_formatted = json.dumps(decoded_access_token, indent=2)
        logger.info("Decoded Access Token:\n%s", access_token_formatted)

        # Token Expiry
        token_expiry = datetime.fromtimestamp(int(decoded_access_token["exp"]))
        logger.info("Token Expires at: %s", str(token_expiry))
        return access_token
    except Exception as err:
        logger.error("Error in msgraph_auth: %s", str(err))
        raise


def make_ms_graph_request(access_token: str, url: str) -> Dict:
    """
    Make a request to the Microsoft Graph API.

    :param access_token: The access token for Microsoft Graph API authentication.
    :param url: The URL for the Microsoft Graph API endpoint.
    :return: The JSON response from the Microsoft Graph API.
    :raises Exception: If there's an HTTP error or other issues in making the request.
    """
    headers = {"Authorization": f"Bearer {access_token}"}
    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        return response.json()
    except requests.exceptions.HTTPError as err:
        logger.error(f"HTTP Error: {err}")
        raise
    except Exception as err:
        logger.error(f"Error in make_ms_graph_request: {err}")
        raise


def get_site_id(site_domain: str, site_name: str, access_token: str) -> Optional[str]:
    """
    Get the Site ID from Microsoft Graph API.

    :param site_domain: The domain of the site in Microsoft Graph.
    :param site_name: The name of the site in Microsoft Graph.
    :param access_token: The access token for Microsoft Graph API authentication.
    :return: The Site ID or None if there's an error.
    """
    endpoint = (
        f"https://graph.microsoft.com/v1.0/sites/{site_domain}:/sites/{site_name}:/"
    )
    try:
        logger.info("Getting the Site ID...")
        result = make_ms_graph_request(access_token, endpoint)
        site_id = result.get("id")
        if site_id:
            logger.info(f"Site ID retrieved: {site_id}")
            return site_id
    except Exception as err:
        logger.error(f"Error retrieving Site ID: {err}")
        return None


def get_drive_id(access_token: str, site_id: str) -> str:
    """
    Get the drive ID from a Microsoft Graph site.

    :param access_token: The access token for Microsoft Graph API authentication.
    :param site_id: The site ID in Microsoft Graph.
    :return: The drive ID.
    :raises Exception: If there's an error in fetching the drive ID.
    """
    url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drive"
    try:
        json_response = make_ms_graph_request(access_token, url)
        return json_response["id"]
    except Exception as err:
        logger.error(f"Error in get_drive_id: {err}")
        raise


def get_files_in_site(
    access_token: str,
    site_id: str,
    drive_id: str,
    minutes_ago: Optional[int] = None,
    file_formats: Optional[List[str]] = None,
) -> List[Dict]:
    """
    Get a list of files in a site's drive, optionally filtered by creation or last modification time and file formats.

    :param access_token: The access token for Microsoft Graph API authentication.
    :param site_id: The site ID in Microsoft Graph.
    :param drive_id: The drive ID in Microsoft Graph.
    :param minutes_ago: Optional integer to filter files created or updated within the specified number of minutes from now. If None, no time-based filtering is applied.
    :param file_formats: List of desired file formats.
    :return: A list of file details.
    :raises Exception: If there's an error in fetching file details.
    """
    url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{drive_id}/root/children"
    try:
        logger.info("Making request to Microsoft Graph API")
        json_response = make_ms_graph_request(access_token, url)
        files = json_response["value"]
        logger.info("Received response from Microsoft Graph API")

        time_limit = (
            datetime.now(timezone.utc) - timedelta(minutes=minutes_ago)
            if minutes_ago is not None
            else None
        )

        filtered_files = [
            file
            for file in files
            if (
                (
                    time_limit is None
                    or datetime.fromisoformat(
                        file["fileSystemInfo"]["createdDateTime"].rstrip("Z")
                    ).replace(tzinfo=timezone.utc)
                    >= time_limit
                    or datetime.fromisoformat(
                        file["fileSystemInfo"]["lastModifiedDateTime"].rstrip("Z")
                    ).replace(tzinfo=timezone.utc)
                    >= time_limit
                )
                and (
                    not file_formats
                    or any(file["name"].endswith(f".{fmt}") for fmt in file_formats)
                )
            )
        ]

        return filtered_files
    except Exception as err:
        logger.error(f"Error in get_files_in_site: {err}")
        raise


def get_file_permissions(access_token: str, site_id: str, item_id: str) -> List[Dict]:
    """
    Get the permissions of a file in a site.

    :param access_token: The access token for Microsoft Graph API authentication.
    :param site_id: The site ID in Microsoft Graph.
    :param item_id: The item ID of the file in Microsoft Graph.
    :return: A list of permission details.
    :raises Exception: If there's an error in fetching permission details.
    """
    url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drive/items/{item_id}/permissions"

    try:
        json_response = make_ms_graph_request(access_token, url)
        return json_response["value"]
    except Exception as err:
        logger.error(f"Error in get_file_permissions: {err}")
        raise


def group_users_by_role(user_data: List[Dict]) -> str:
    """
    Group users by their roles and return a JSON string representing this grouping.

    This function takes a list of user data dictionaries, each containing user roles, and other attributes.
    It returns a JSON string where the keys are roles and the values are lists of user display names corresponding to each role.

    :param user_data: List of dictionaries containing user data.
    :return: JSON string representing users grouped by their roles.
    :raises DataError: If there are issues with data processing.
    """
    try:
        grouped_users = {}
        for user in user_data:
            roles = user.get('roles', [])
            display_name = user.get('grantedTo', {}).get('user', {}).get('displayName') or \
                           user.get('grantedToV2', {}).get('siteGroup', {}).get('displayName') or \
                           user.get('grantedToV2', {}).get('group', {}).get('displayName')

            for role in roles:
                if role in grouped_users:
                    grouped_users[role].append(display_name)
                else:
                    grouped_users[role] = [display_name]

        return json.dumps(grouped_users)
    except Exception as e:
        logger.error(f"Error processing user data: {e}")


def get_docx_content(
    site_id: str, file_name: str, drive_id: str, access_token: str
) -> Optional[str]:
    """
    Retrieve the text content from a .docx file in a specific site drive.

    :param site_id: The site ID in Microsoft Graph.
    :param file_name: The name of the .docx file.
    :param drive_id: The drive ID in Microsoft Graph.
    :param access_token: The access token for Microsoft Graph API authentication.
    :return: Text content of the docx file or None if there's an error.
    """

    endpoint = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{drive_id}/root:/{file_name}:/content"
    try:
        logger.info(f"Starting request for file: {file_name}")
        response = requests.get(
            endpoint, headers={"Authorization": "Bearer " + access_token}
        )

        # Log different outcomes of the response
        if response.status_code != 200:
            logger.error(
                f"Failed to retrieve file content. Status code: {response.status_code}, Response: {response.text}"
            )
            return None

        file_content = response.content
        document = DocxDocument(io.BytesIO(file_content))
        content = "\n".join([paragraph.text for paragraph in document.paragraphs])

        logger.info(f"Successfully retrieved content for file: {file_name}")
        return content
    except requests.exceptions.RequestException as req_err:
        logger.error(f"Request error retrieving content for {file_name}: {req_err}")
        return None
    except Exception as err:
        logger.error(f"General error retrieving content for {file_name}: {err}")
        return None


def retrieve_sharepoint_files_content(site_domain: str, site_name: str, minutes_ago: Optional[int] = None, file_formats: Optional[List[str]] = None) -> Dict[str, Dict[str, Optional[str]]]:
    """
    Retrieve contents of files from a specified SharePoint location, optionally filtering by last modification time and file formats.

    :param site_domain: The domain of the site in Microsoft Graph.
    :param site_name: The name of the site in Microsoft Graph.
    :param minutes_ago: Optional; filter for files modified within the specified number of minutes.
    :param file_formats: Optional; list of desired file formats to include.
    :return: Dictionary with file names as keys and a dictionary containing their content, location, and users_by_role as values.
    """
    try:
        client_auth = msgraph_auth(client_id=clientID, client_secret=clientSecret, authority=authority, scope=scope)
        access_token = client_auth["access_token"]

        site_id = get_site_id(site_domain, site_name, access_token)
        if not site_id:
            logger.error("Failed to retrieve site_id")
            return None

        drive_id = get_drive_id(access_token, site_id)
        if not drive_id:
            logger.error("Failed to retrieve drive ID")
            return None

        files = get_files_in_site(access_token, site_id, drive_id, minutes_ago, file_formats)
        if not files:
            logger.error("No files found in the site's drive")
            return None

        file_contents = []
        for file in files:
            file_name = file.get("name")
            if file_name and (not file_formats or any(file_name.endswith(f'.{fmt}') for fmt in file_formats)):
                logger.info(f"Fetching content for file: {file_name}")
                content = get_docx_content(site_id, file_name, drive_id, access_token)
                url_location = f'https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{drive_id}/root:/{file_name}:/content'
                users_by_role = group_users_by_role(get_file_permissions(access_token, site_id, file["id"]))
                sec_group = manager_security.get_highest_priority_group(users_by_role)
                print(f"sec_group: {sec_group}")
                file_content = Document(page_content=content, metadata={"source": url_location
                    ,"read_access_group": users_by_role,
                    "security_group": str(sec_group)})
                file_contents.append(file_content)
            else:
                logger.info(f"Skipping file: {file_name}")

        return file_contents
    except Exception as e:
        logger.error(f"Error in retrieve_sharepoint_files_content function: {e}")
        return None